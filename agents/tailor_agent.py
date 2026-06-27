"""
agents/tailor_agent.py
Stage 3 — Resume Tailor

Builds a one-page, job-specific resume from the master resume:
- Preserves name and contact header exactly
- Rewrites summary and reorders/filters bullets for relevance
- Drops irrelevant projects, skills, and sections
"""

import re
from pathlib import Path
from datetime import datetime

from rich.console import Console
from docx import Document
from docx.shared import Pt, Inches

from core.config import OUTPUT_DIR, MASTER_RESUME_PATH, PROFILE
from core.database import get_jobs_by_status, update_job_status
from core.resume_parser import load_resume_text
from core.llm import ask
from core.text_utils import sanitize_plain_text, sanitize_llm_object

console = Console()

RESUMES_DIR = OUTPUT_DIR / "resumes"
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

TAILOR_SYSTEM = """You are an expert resume writer who tailors resumes for specific job postings.
You must keep every claim factual — rephrase and reorder only, never invent experience.
Your output must fit on ONE printed page: be selective, concise, and aggressive about removing irrelevant content.
Use plain ASCII punctuation only: regular hyphen (-), straight quotes, commas, and periods. No em dashes, en dashes, or smart quotes.
Respond in valid JSON only."""

TAILOR_PROMPT = """
Tailor this candidate's resume for the target job below.

## MASTER RESUME
{resume_text}

## TARGET JOB
Title: {title}
Company: {company}
Description:
{description}

## HARD RULES
1. ONE PAGE ONLY - the final resume must fit a single page when printed.
2. Keep the candidate's name and contact info unchanged (do not include them in JSON — they are added separately).
3. Write exactly ONE professional summary — 2 sentences max, focused on this role.
4. Include ONLY experience bullets that are relevant to this job. Drop irrelevant bullets entirely.
   - Max 5 bullets for the main role.
   - Rewrite bullets to emphasize matching keywords from the job description.
5. Include ONLY projects directly relevant to this job (max 2 one-line entries). Omit the projects section entirely if none apply.
6. Include ONLY skill lines relevant to this job (max 4 compact lines). Remove unrelated languages, frameworks, and tools.
7. Keep education compact: degree line, institution, and at most one extra highlight line.
8. Include "additional" entries only if they add clear value for THIS job (max 2). Otherwise return an empty list.
9. Do NOT duplicate the summary anywhere else. Do NOT add a second opening paragraph.

Return this JSON structure:
{{
  "summary": "<2 sentences max, tailored to the job>",
  "experience": [
    {{
      "title_line": "<Job Title>\\t<dates>",
      "company": "<Company Name>",
      "bullets": ["<relevant bullet>", "..."]
    }}
  ],
  "projects": [
    {{"line": "<Project Name> - <one concise line>"}}
  ],
  "skills": [
    "<Category: item1, item2, item3>"
  ],
  "education": {{
    "degree_line": "<degree and year>",
    "institution": "<school name>",
    "extra_line": "<optional one line or empty string>"
  }},
  "additional": ["<optional bullet>"],
  "languages_line": "Hebrew - Native  |  English - Fluent",
  "tailoring_notes": "<brief note on what you kept, cut, and emphasized>"
}}
"""


def _safe_filename(text: str) -> str:
    clean = re.sub(r"[^\w\s-]", "", text).strip()
    clean = re.sub(r"[\s]+", "_", clean)
    return clean[:60]


def _read_master_header(master_path: Path) -> tuple[str, str]:
    doc = Document(str(master_path))
    name = doc.paragraphs[0].text.strip() if doc.paragraphs else PROFILE.get("name", "")
    contact = doc.paragraphs[1].text.strip() if len(doc.paragraphs) > 1 else ""
    if not contact:
        contact = (
            f"{PROFILE.get('location', '')}  |  {PROFILE.get('phone', '')}  |  "
            f"{PROFILE.get('email', '')}\nGitHub: {PROFILE.get('github', '')}  |  "
            f"LinkedIn: {PROFILE.get('linkedin', '')}"
        )
    return name, contact


def _set_compact_page(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)


def _style_paragraph(paragraph, *, bold: bool = False, size: int = 10) -> None:
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.bold = bold


def _add_text(doc: Document, text: str, *, bold: bool = False, bullet: bool = False) -> None:
    if not text or not text.strip():
        return
    text = sanitize_plain_text(text.strip())
    style = "List Bullet" if bullet else None
    para = doc.add_paragraph(text, style=style)
    _style_paragraph(para, bold=bold)


def tailor_resume_content(job: dict, resume_text: str) -> dict:
    """Ask Claude to tailor the resume for this job. Returns structured JSON."""
    prompt = TAILOR_PROMPT.format(
        resume_text=resume_text[:6000],
        title=job["title"],
        company=job["company"] or "Unknown",
        description=(job["description"] or "")[:4000],
    )
    return sanitize_llm_object(
        ask(prompt, system=TAILOR_SYSTEM, max_tokens=2500, json_mode=True)
    )


def build_docx(original_path: Path, tailored: dict, job: dict) -> Path:
    """Build a clean one-page tailored resume. Preserves header from master."""
    name, contact = _read_master_header(original_path)

    doc = Document()
    _set_compact_page(doc)

    # Header — never modified by the LLM
    name_para = doc.add_paragraph(name)
    _style_paragraph(name_para, bold=True, size=14)
    _add_text(doc, contact)

    # Summary — single block only
    _add_text(doc, "PROFESSIONAL SUMMARY", bold=True)
    _add_text(doc, tailored.get("summary", ""))

    # Experience
    experience = tailored.get("experience") or []
    if experience:
        _add_text(doc, "PROFESSIONAL EXPERIENCE", bold=True)
        for role in experience:
            _add_text(doc, role.get("title_line", ""), bold=True)
            _add_text(doc, role.get("company", ""))
            for bullet in role.get("bullets") or []:
                _add_text(doc, bullet, bullet=True)

    # Projects — only if Claude included relevant ones
    projects = tailored.get("projects") or []
    project_lines = [p.get("line", "") for p in projects if p.get("line")]
    if project_lines:
        _add_text(doc, "PROJECTS", bold=True)
        for line in project_lines[:2]:
            _add_text(doc, line)

    # Education
    education = tailored.get("education") or {}
    if education:
        _add_text(doc, "EDUCATION", bold=True)
        _add_text(doc, education.get("degree_line", ""))
        _add_text(doc, education.get("institution", ""))
        _add_text(doc, education.get("extra_line", ""))

    # Skills
    skills = [s for s in (tailored.get("skills") or []) if s.strip()]
    if skills:
        _add_text(doc, "TECHNICAL SKILLS", bold=True)
        for skill_line in skills[:4]:
            _add_text(doc, skill_line, bullet=True)

    # Languages
    languages_line = tailored.get("languages_line") or "Hebrew - Native  |  English - Fluent"
    _add_text(doc, "LANGUAGES", bold=True)
    _add_text(doc, languages_line)

    # Additional — only if relevant
    additional = [a for a in (tailored.get("additional") or []) if a.strip()]
    if additional:
        _add_text(doc, "ADDITIONAL", bold=True)
        for item in additional[:2]:
            _add_text(doc, item, bullet=True)

    fname = _safe_filename(f"{job['title']}_{job['company'] or 'company'}")
    timestamp = datetime.now().strftime("%Y%m%d")
    out_path = RESUMES_DIR / f"resume_{fname}_{timestamp}.docx"
    doc.save(str(out_path))
    return out_path


def tailor_job(job: dict, resume_text: str = None) -> Path:
    """Tailor and save a resume for one job. Returns the output file path."""
    resume_text = resume_text or load_resume_text()
    tailored = tailor_resume_content(job, resume_text)
    out_path = build_docx(MASTER_RESUME_PATH, tailored, job)
    update_job_status(
        job["job_id"],
        status="tailored",
        tailored_resume_path=str(out_path),
    )
    return out_path


def run() -> dict:
    """Tailor resumes for all 'approved' jobs."""
    console.rule("[bold blue]✍️  Resume Tailor — Stage 3")

    approved_jobs = get_jobs_by_status("approved")
    if not approved_jobs:
        console.print("[yellow]No approved jobs to tailor. Run the scorer first.[/yellow]")
        return {"tailored": 0}

    console.print(f"Tailoring resume for [bold]{len(approved_jobs)}[/bold] approved jobs...\n")
    resume_text = load_resume_text()
    stats = {"tailored": 0, "failed": 0}

    for job in approved_jobs:
        console.print(f"  ✍️  {job['title']} @ {job['company'] or 'N/A'} (score: {job['match_score']})")
        try:
            out_path = tailor_job(job, resume_text=resume_text)
            console.print(f"     [green]✅ Saved:[/green] {out_path.name}")
            stats["tailored"] += 1
        except Exception as e:
            console.print(f"     [red]❌ Failed: {e}[/red]")
            stats["failed"] += 1

    console.rule("[bold green]Tailoring Complete")
    console.print(f"  Tailored: [bold green]{stats['tailored']}[/bold green] | Failed: {stats['failed']}")
    console.print(f"  📁 Resumes saved to: [cyan]{RESUMES_DIR}[/cyan]")
    return stats


if __name__ == "__main__":
    run()
