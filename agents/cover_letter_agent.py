"""
agents/cover_letter_agent.py
Stage 4 — Cover Letter Writer

For each tailored job, Claude writes a personalized cover letter
connecting your specific experience to the job's exact requirements.
Outputs both .docx and .txt per job.
"""

import re
from pathlib import Path
from datetime import datetime

from rich.console import Console
from docx import Document

from core.config import OUTPUT_DIR, PROFILE
from core.database import get_jobs_by_status, update_job_status
from core.resume_parser import load_resume_text
from core.llm import ask
from core.text_utils import sanitize_plain_text

console = Console()

LETTERS_DIR = OUTPUT_DIR / "cover_letters"
LETTERS_DIR.mkdir(parents=True, exist_ok=True)

LETTER_SYSTEM = """You are an expert career coach who writes compelling, authentic cover letters.
Your letters are specific, warm, and professional - never generic.
You connect the candidate's actual experience to the job's real requirements.
Use plain ASCII punctuation only: regular hyphen (-), straight quotes, commas, and periods. No em dashes or smart quotes."""

LETTER_PROMPT = """
Write a cover letter for this job application.

## CANDIDATE
Name: {name}
Email: {email}
Resume Summary:
{resume_text}

## JOB
Title: {title}
Company: {company}
Description:
{description}

## INSTRUCTIONS
- Opening: Hook with why this specific company/role excites them (be specific, not generic)
- Body paragraph 1: Highlight the 2-3 most relevant experiences from their background
- Body paragraph 2: Address a specific challenge/need in the job description and how they solve it
- Closing: Clear call to action, confident but not arrogant
- Tone: Professional but human — avoid corporate buzzwords
- Length: 3-4 paragraphs, max 350 words
- Do NOT start with "I am writing to apply for..."

Return ONLY the cover letter text, no JSON, no extra commentary.
"""


def write_cover_letter(job: dict, resume_text: str) -> str:
    prompt = LETTER_PROMPT.format(
        name=PROFILE.get("name", ""),
        email=PROFILE.get("email", ""),
        resume_text=resume_text[:2000],
        title=job["title"],
        company=job["company"] or "your company",
        description=(job["description"] or "")[:2500],
    )
    return sanitize_plain_text(
        ask(prompt, system=LETTER_SYSTEM, max_tokens=800)
    )


def save_cover_letter(text: str, job: dict) -> Path:
    clean = re.sub(r'[^\w\s-]', '', f"{job['title']}_{job['company'] or 'co'}").strip()
    clean = re.sub(r'\s+', '_', clean)[:50]
    timestamp = datetime.now().strftime("%Y%m%d")

    # Save as .docx
    doc = Document()
    text = sanitize_plain_text(text)
    doc.add_paragraph(PROFILE.get("name", ""))
    doc.add_paragraph(f"{PROFILE.get('email', '')} | {PROFILE.get('location', '')}")
    doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph(f"Hiring Team - {job['company'] or ''}")
    doc.add_paragraph()
    for paragraph in text.strip().split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(sanitize_plain_text(paragraph.strip()))
    doc.add_paragraph()
    doc.add_paragraph(f"Best regards,\n{PROFILE.get('name', '')}")

    out_path = LETTERS_DIR / f"cover_{clean}_{timestamp}.docx"
    doc.save(str(out_path))

    # Also save plain text
    txt_path = LETTERS_DIR / f"cover_{clean}_{timestamp}.txt"
    txt_path.write_text(text)

    return out_path


def run() -> dict:
    console.rule("[bold blue]📝 Cover Letter Writer — Stage 4")

    jobs = get_jobs_by_status("tailored")
    if not jobs:
        console.print("[yellow]No tailored jobs found. Run the tailor agent first.[/yellow]")
        return {"written": 0}

    console.print(f"Writing cover letters for [bold]{len(jobs)}[/bold] jobs...\n")
    resume_text = load_resume_text()
    stats = {"written": 0, "failed": 0}

    for job in jobs:
        console.print(f"  📝 {job['title']} @ {job['company'] or 'N/A'}")
        try:
            letter_text = write_cover_letter(job, resume_text)
            out_path = save_cover_letter(letter_text, job)
            update_job_status(
                job["job_id"],
                status="ready",
                cover_letter_path=str(out_path)
            )
            console.print(f"     [green]✅ Saved:[/green] {out_path.name}")
            stats["written"] += 1
        except Exception as e:
            console.print(f"     [red]❌ Failed: {e}[/red]")
            stats["failed"] += 1

    console.rule("[bold green]Cover Letters Complete")
    console.print(f"  Written: [bold green]{stats['written']}[/bold green]")
    console.print(f"  📁 Saved to: [cyan]{LETTERS_DIR}[/cyan]")
    return stats


if __name__ == "__main__":
    run()
